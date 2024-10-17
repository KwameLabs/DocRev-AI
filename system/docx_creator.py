import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_word_doc(contents, filename):
    # Create a document
    doc = docx.Document()

    # Add a paragraph to the document
    p = doc.add_paragraph()

    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0

    # Add a run to the paragraph
    run = p.add_run("NATIONAL DEVELOPMENT PLANNING COMMISSION")

    # Add some formatting to the run
    
    run.bold = True
    run.italic = True
    run.underline = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(14)

#    Add another paragraph (left blank for an empty line)
    doc.add_paragraph()
    # Add more text to the same paragraph
    
    
    run = p.add_run("PUBLIC POLICY REVIEW ADVISORY")

    # Format the run
    run.bold = True
    run.underline = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(14)

    # Add another paragraph (left blank for an empty line)
    doc.add_paragraph()
    #doc.add_paragraph()
    
    # Add another paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Add a run and format it
    run = p.add_run(contents)
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)


    # Save the document
    
    doc.save(filename+".docx")