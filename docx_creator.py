import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def get_download_folder():
        home = os.path.expanduser("~")
        download_folder = os.path.join(home, "Downloads\\")
        # Normalize the path to ensure it’s interpreted correctly
        return download_folder
    
def create_word_doc(contents, filename):
    # Create a document
    doc = docx.Document()

    # Add a paragraph to the document
    p = doc.add_paragraph()

    # Add some formatting to the paragraph
    p.paragraph_format.line_spacing = 1
    p.paragraph_format.space_after = 0

    # Add a run to the paragraph
    run = p.add_run("NATIONAL DEVELOPMENT PLANNING COMMISSION")

    # Add some formatting to the run
    
    run.bold = True
    # run.italic = True
    run.underline = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(14)

#    Add another paragraph (left blank for an empty line)
    p=doc.add_paragraph()
    # Add more text to the same paragraph
    p.paragraph_format.line_spacing = 2
    p.paragraph_format.space_after = 0
    
    
    
    run = p.add_run("PUBLIC POLICY REVIEW ADVISORY")

    # Format the run
    run.bold = True
    run.underline = True
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(14)

    # Add another paragraph (left blank for an empty line)
    doc.add_paragraph()
    #doc.add_paragraph()
    
    # Add another paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Add a run and format it
    run = p.add_run(contents)
    run.font.name = 'Arial'
    run.font.size = docx.shared.Pt(12)
    
     # Add another paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT



    # Save the document
    # complete_path = get_download_folder()+filename+".docx"
    filename_without_extension = '.'.join(filename.split('.')[:-1])
    complete_path = os.path.join(get_download_folder(), filename_without_extension + ".docx")

    
    try:
        doc.save(complete_path)
        print(f"Document saved successfully at {complete_path}")
    except PermissionError:
        print(f"Error: Unable to save the file. Please ensure that {complete_path} is not open in another application.")
   
    